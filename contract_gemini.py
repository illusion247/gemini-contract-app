import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
import io
import PyPDF2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
import base64

load_dotenv()

# Set up Gemini API client
api_key = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=api_key)

# Load the model
model = genai.GenerativeModel("gemini-1.5-flash")  # Use gemini-1.5-flash

# Function to extract info directly from PDF with Gemini
def extract_info_gemini_vision(pdf_file):
    if pdf_file:
        prompt = f"""
            Analyze the following contract document and extract the following information:

            1. Termination Notice No. of Days:  How many days are required to give for a termination notice, and indicate which party is terminating the contract.
            2. Auto Renewal: Does the contract contains a renewal clause, if so please include details. Find infromation that refers to the renewal of contract.
            3. Signed Date of the Client (client): Extract the date signed by the client.
            4. Effectivity Date: find infromation or clause about the effectivity date of the contract.

            additional instruction: provide the page number and section for reference if information is available

            Note: The Service Provider is always Towers Watson or Willis Towers Watson. Please extract only the Client's Signature Date.
            """
        try:
             pdf_content = pdf_file.read()
             response = model.generate_content([prompt, {"mime_type": "application/pdf", "data": pdf_content}])
             return response.text
        except Exception as e:
            return f"Error querying Gemini API: {e}"
    else:
         return "No file Uploaded"


def create_word_document(extracted_data, pdf_file, pdf_file_name):
    document = Document()
    document.styles['Normal'].font.size = Pt(12)
    section = document.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    document.add_heading(f"Contract Analysis for {pdf_file_name}", level=1)
    if extracted_data:
        document.add_paragraph("Extracted Information:", style="Heading 2")
        document.add_paragraph(extracted_data)
    else:
        document.add_paragraph("No information extracted.")

    #highlight keywords
    try:
       pdf_text = ""
       pdf_file_object = io.BytesIO(pdf_file.read())
       pdf_reader = PyPDF2.PdfReader(pdf_file_object)
       for page_num in range(len(pdf_reader.pages)):
           page = pdf_reader.pages[page_num]
           page_text = page.extract_text()
           pdf_text+=page_text

       keywords = ["termination", "renewal", "date", "effectivity"]

       for keyword in keywords:
           for index in range(0, len(pdf_text), 100):
               chunk = pdf_text[index:index + 100] # split into 100 character chunks
               if keyword in chunk.lower():
                p = document.add_paragraph()
                p.add_run(chunk.strip()).font.highlight_color = RGBColor(255, 255, 0)  #Highlight it yellow
                p.add_comment(keyword)
    except Exception as e:
           document.add_paragraph(f"Error highlighting the words: {e}")

    doc_bytes = io.BytesIO()
    document.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()

# Streamlit app
st.title("Contract Information Extractor")
uploaded_file = st.file_uploader("Upload a Contract PDF", type="pdf")

if uploaded_file:
    with st.spinner("Extracting information..."):
       extracted_data = extract_info_gemini_vision(uploaded_file)

       if extracted_data:
          st.success("Information extracted successfully!")
          st.write(extracted_data)
          word_file = create_word_document(extracted_data, uploaded_file, uploaded_file.name)
          st.download_button(
               label="Export as Word Document",
               data = word_file,
               file_name=f"{uploaded_file.name}.docx",
               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
             )

       else:
          st.error("Error during information extraction.")
